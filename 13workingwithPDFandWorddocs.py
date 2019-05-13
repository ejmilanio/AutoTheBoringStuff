###---------------------------------------------###
#      Ch. 13 Working with PDF and Word Docs      #
###---------------------------------------------###

import os
os.chdir("C:\\Users\\ejmil\\source\\repos\\AutoTheBoringStuff")

### Extracting Text from PDF files ###
import PyPDF2 
pdfFileObj = open('meetingminutes.pdf', 'rb')           #open PDF file in 'read binary' mode; creates a PDFFile Object
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)            #create a PdfFileReader object and pass it PDFFile obj

pdfReader.numPages                                      #attr of a PdfFileReader object
pageObj = pdfReader.getPage(0)                          #.getPage() takes a page. An arrray so starts from 0
pageObj.extractText()                                   #.extractText() function extracts text

### Decrypting PDFs ###
#some PDFs are encrypted and need passwords

import PyPDF2
pdfReader = PyPDF2.PdfFileReader(open('encrypted.pdf', 'rb'))
pdfReader.isEncrypted                                   #checks if encrypted; boolean value
pdfReader.decrypt('rosebud')                            #returns 1 if password is correct; 0 if password is incorrect
pdfReader.getPage(0)

### Creating PDFs ###
#Must create a PdfFileWriter object to create a PDF
#PyPDF2 can only copy pages from other PDFs, roatate pages, overlay pages and encrypt files

#PyPDF2 cant directly edit a PDF. Must create new PDF and copy content to the new PDF
    #1. Open one ore more existing PDFs (the source PDFs) into PdfFileReader objects
    #2. Create a new PdfFileWriter Object.
    #3. Copy pages from the PdfFileReader objects into the PdfFileWriter Object
    #4. Finally, use the PdfFileWriter object to write the output PDF.

#PdfFileWriter object only represents a PDF doc and doesn't actually create one
    #PdfFileWriter's write() method will create a PDF file

#@ Copying Pages @#

import PyPDF2
pdf1File = open('meetingminutes.pdf', 'rb')         #open both PDFs in 'read binary' mode
pdf2File = open('meetingminutes2.pdf', 'rb')
pdf1Reader = PyPDF2.PdfFileReader(pdf1File)         #create a PdfFileReader object for both PDFs
pdf2Reader = PyPDF2.PdfFileReader(pdf2File)
pdfWriter = PyPDF2.PdfFileWriter()                  #create a PdfFileWriter for new combined PDF

for pageNum in range(pdf1Reader.numPages):          #loop through each page in the 1st PDF
    pageObj = pdf1Reader.getPage(pageNum)           #create a Page object for each page
    pdfWriter.addPage(pageObj)                      #append each Page object to the PdfFileWriter object

for pageNum in range(pdf2Reader.numPages):          #Loop through each page in the 2nd PDF and append to the PdfFileWriter
    pageObj = pdf2Reader.getPage(pageNum)           #create a Page object for each page
    pdfWriter.addPage(pageObj)                      #append each Page object to the PdfFileWriter object

pdfOutputFile = open('combinedminutes.pdf', 'wb')   #create a new PDF by passing a File obj to the PdfFileWriters write() method
pdfWriter.write(pdfOutputFile)
pdfOutputFile.close()
pdf1File.close()
pdf2File.close()

#@ Rotating Pages @#
#Pages can be rotated 90 degree increments with the rotateClockwise() and rotateCounterClockwise() methods
    #can be passed 90, 180 or 270 arguments 

import PyPDF2
minutesFile = open('meetingminutes.pdf', 'rb')          #open a PDF in 'read binary' mode
pdfReader = PyPDF2.PdfFileReader(minutesFile)           #create a PdfFileReader object
page = pdfReader.getPage(0)                             #create a Page object
page.rotateClockwise(90)                                #rotate that Page object clockwise 90 degrees
pdfWriter = PyPDF2.PdfFileWriter()                      #create a PdfFileWriter object
pdfWriter.addPage(page)                                 #append roated page to the object
resultPdfFile = open('rotatedPage.pdf', 'wb')           #open the PdfFileWriter object
pdfWriter.write(resultPdfFile)                          #save the PdfFilewriter object
resultPdfFile.close()                                   #close the new PdfFileWriter object
minutesFile.close()                                     #close original PdfFileWriter object

#@ OVerlaying Pages @#
#good for adding things like watermarks or logos

import PyPDF2
minutesFile = open('meetingminutes.pdf', 'rb')                          #create File obj by opening PDF file
pdfReader = PyPDF2.PdfFileReader(minutesFile)                           #create a PdfReader obj
minutesFirstPage = pdfReader.getPage(0)                                 #create a Page obj
pdfWatermarkReader = PyPDF2.PdfFileReader(open('watermark.pdf','rb'))   #create a PdfReader obj of page to be laid over
minutesFirstPage.mergePage(pdfWatermarkReader.getPage(0))               #merge watermaker PdfReader obj over first PdfReader obj
pdfWriter = PyPDF2.PdfFileWriter()                                      #Create a PdfFileWriter obj
pdfWriter.addPage(minutesFirstPage)                                     #add new watermarked cover to new PDF

for pageNum in range(1,pdfReader.numPages):                             #loop through remaining pages in original doc
    pageObj = pdfReader.getPage(pageNum)                                    #and add to new PDF doc
    pdfWriter.addPage(pageObj)
resultPdfFile = open('watermarkedCover.pdf', 'wb')                      #create, save and close final PDF doc
pdfWriter.write(resultPdfFile)  
minutesFile.close()
resultPdfFile.close()

#@ Encrypting PDFs @#
#PdfFileWriter Obj can also add encryptions
import PyPDF2
pdfFile = open('meetingminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFile)
pdfWriter = PyPDF2.PdfFileWriter()
for pageNum in range(pdfReader.numPages):
    pdfWriter.addPage(pdfReader.getPage(pageNum))

pdfWriter.encrypt('swordfish')                              #call a PdfFileWriter obj encrypt function and pass a string as a password to encrypt
resultPdf = open('encryptedminutes.pdf', 'wb')
pdfWriter.write(resultPdf)
resultPdf.close()

### Word Documents ###
#.docx files have three different data types represented in Python-docx
    #Document object represents the entire document
    #Paragraph object represents a paragraph
    #Run object represents a contiguous run of text w/ the same style

#@ Reading Word Docs @#
import docx
doc = docx.Document('demo.docx')            #docx.Document('documentName.docx') will create a Document object
len(doc.paragraphs)                         #.paragraphs is an attribute of a Document object and creates a list of Paragraph obj when called
doc.paragraphs[0].text                      #.text is an attr of a Paragraph obj
doc.paragraphs[1].text                      #can access different Paragraph obj by calling it's pos in the array
len(doc.paragraphs[1].runs)                 #.runs is an attr of Paragraph obj. Creates a list of Run obj
doc.paragraphs[1].runs[0].text              #Run objs have a text attr as well.
doc.paragraphs[1].runs[2].text
doc.paragraphs[1].runs[3].text

#@ Getting the Full Text from a .docx File @#
import readDocx
print(readDocx.getText('demo.docx'))

import docx
doc = docx.Document('demo.docx')
fullText = []
for para in doc.paragraphs:
    fullText.append('    ' + para.text)
print('\n\n'.join(fullText))

#@ Styling Paragraph and Run Objects @#
# Types of stylings
    #Paragraph styles which can only be applied to Paragraph objects
    #Run styles which can only be applied to Run objects
    #Linked styles which can be applied to both
#use the .style attr for either Paragraph obj or Run Obj
    # paragraphObj.style = 'style'
    
''' Types of styles in word

'Normal'    'Heading5'  'ListBullet'    'ListParagraph'     'BodyText'

'Heading6'  'ListBullet2'   'MacroText'     'BodyText2'     'Heading7'

'ListBullet3'   'NoSpacing'     'BodyText3'     'Heading8'  'ListContinue'

'Quote'     'Caption'   'Heading9'  'ListContinue2'

'Subtitle'  'Heading1'  'IntenseQuote'  'ListContinue3'

'TOCHeading'    'Heading2'     'List'   'ListNumber'

'Title'     'Heading3'     'List2'  'ListNumber2'

 'Heading4'      'List3'     'ListNumber3'

'''

#@ Creating Word Documents with NonDefault Styles @#

#have to create the new style in Word
    #at the bottom of the styles pane click New Style
    #the name given to this style will be available for use with the Python-Docx module

#@ Run Attributes @#
#Runs can be styled using text attributes
    #Three states for each attr
        #True, attr is enabled
        #False, attr is disabled
        #None, text is set to Run Objs default setting

""".text Attributes

Attribute           Description
bold                Text appears bold
italic              Text appears italicized
underline           Text is underlined
strike              Text has a strikethrough
double_strike       Text has a double strikethrough
all_caps            Text is all caps
small_caps          The text appears in capital letters, w/ lowercase letters appearing 2 points smaller
shadow              Text appears with a shadow
outline             Text appears with an outline rather than solid
rtl                 Text is written right to left
imprint             The text appears pressed into the page
emboss              The text appears rased off the page in releif


"""

import docx
doc = docx.Document('demo.docx')
doc.paragraphs[0].text
doc.paragraphs[0].style
doc.paragraphs[0].style = 'Normal'
doc.paragraphs[1].text
(doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text)
doc.paragraphs[1].runs[0].style = 'QuoteChar'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save('restyled.docx')

#@ Writing Word Documents @#

import docx
doc = docx.Document()
doc.add_paragraph('Hello world!')           #Document object's .add_paragraph method adds a paragraph
doc.save('helloworld.docx')


import docx
doc = docx.Document()
doc.add_paragraph('Hello World!')
paraObj1 = doc.add_paragraph('This is a second paragraph')
paraObj2 = doc.add_paragraph('This is a third paragraph')
paraObj1.add_run(' This text is being added to the second paragraph.')      #add_run() will add a run
doc.save('multipleParagraphs.docx')


#
import docx
doc = docx.Document()
doc.add_heading('Header 0', 0)              #add_header('text', header num) will add a header of the specified style
doc.add_heading('Header 1', 1)              #creates a Paragraph object
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)
doc.save('heading.docx')

import docx
doc = docx.Document()
doc.add_paragraph('This is on the first page!')
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)         #docx.enum.text.WD_BREAK.PAGE will create a page break at specified run
doc.add_paragraph('This is on the second page!')
doc.save('twoPage.docx')

import docx
doc.add_picture('zophie.png', width= docx.shared.Inches(1), height= docx.shared.Cm(4))
#first argument is filename, optional: width and height arguments. Will default to original size of graphic
    #can specify units as well

